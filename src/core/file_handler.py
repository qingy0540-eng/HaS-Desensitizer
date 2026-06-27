"""文件处理模块"""
from pathlib import Path
from typing import Optional


try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class FileHandler:
    """文件读写处理器"""

    SUPPORTED_READ = {
        ".txt": "text",
        ".md": "text",
        ".csv": "text",
        ".json": "text",
        ".xml": "text",
        ".py": "text",
        ".js": "text",
        ".java": "text",
        ".c": "text",
        ".cpp": "text",
        ".h": "text",
        ".go": "text",
        ".rs": "text",
        ".ts": "text",
        ".html": "text",
        ".css": "text",
        ".sql": "text",
        ".log": "text",
    }

    SUPPORTED_WRITE = {
        ".txt": "text",
        ".md": "text",
    }

    @classmethod
    def read(cls, file_path: str | Path, encoding: str = "utf-8") -> str:
        """读取文件内容"""
        path = Path(file_path).resolve()

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        if not path.is_file():
            raise ValueError(f"路径不是文件: {path}")

        suffix = path.suffix.lower()

        if suffix == ".docx" and HAS_DOCX:
            return cls._read_docx(path)
        elif suffix in cls.SUPPORTED_READ:
            return cls._read_text(path, encoding)
        else:
            # 尝试按文本读取
            try:
                return cls._read_text(path, encoding)
            except UnicodeDecodeError:
                raise ValueError(f"不支持的文件格式或编码: {suffix}")

    @classmethod
    def write(
        cls,
        file_path: str | Path,
        content: str,
        encoding: str = "utf-8",
    ) -> None:
        """写入文件"""
        path = Path(file_path).resolve()
        suffix = path.suffix.lower()

        if suffix == ".docx" and HAS_DOCX:
            cls._write_docx(path, content)
        else:
            cls._write_text(path, content, encoding)

    @classmethod
    def _read_text(cls, path: Path, encoding: str) -> str:
        with open(path, "r", encoding=encoding, errors="replace") as f:
            return f.read()

    @classmethod
    def _write_text(cls, path: Path, content: str, encoding: str) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding=encoding) as f:
            f.write(content)

    @classmethod
    def _read_docx(cls, path: Path) -> str:
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @classmethod
    def _write_docx(cls, path: Path, content: str) -> None:
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        doc.save(path)

    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        """获取支持的文件扩展名列表"""
        exts = list(cls.SUPPORTED_READ.keys())
        if HAS_DOCX:
            exts.append(".docx")
        return sorted(exts)

    @classmethod
    def suggest_output_path(cls, input_path: str | Path, suffix: str = "_desensitized") -> Path:
        """建议输出文件路径"""
        path = Path(input_path)
        stem = path.stem
        if not stem.endswith(suffix):
            stem = stem + suffix
        return path.with_name(f"{stem}{path.suffix}")
